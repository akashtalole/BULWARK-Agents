"""Text extraction for uploaded artifact files (PDF/DOCX/TXT).

This runs *before* anything touches the agent pipeline -- it turns
arbitrary file bytes into the plain text `submit_artifact` has always
taken, so Model Armor's injection scan (on Intake/Contract Intelligence's
before_model_callback) still sees every character a judge's PDF
actually contained. Extraction itself is a dumb text pull, not a tool
call, so it needs no agent identity grant.
"""

from __future__ import annotations

import io

from fastapi import UploadFile

MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MiB -- plenty for a compliance doc, not a vector for abuse
_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileError(Exception):
    pass


class EmptyExtractionError(Exception):
    pass


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extension_of(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


async def extract_text(file: UploadFile) -> str:
    """Read `file`, return its extracted plain text. Raises
    UnsupportedFileError for an unrecognized extension, EmptyExtractionError
    if extraction produced nothing usable (e.g. a scanned/image-only PDF
    with no text layer -- this build does no OCR)."""
    ext = extension_of(file.filename)
    if ext not in _ALLOWED_EXTENSIONS:
        raise UnsupportedFileError(
            f"unsupported file type {ext or '(none)'!r} -- only .pdf, .docx, .txt are accepted"
        )

    data = await file.read()
    if len(data) > MAX_UPLOAD_BYTES:
        raise UnsupportedFileError(f"file too large ({len(data)} bytes, max {MAX_UPLOAD_BYTES})")

    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")

    text = text.strip()
    if not text:
        raise EmptyExtractionError(
            "no extractable text found -- if this is a scanned/image PDF, it has no text layer to read"
        )
    return text
